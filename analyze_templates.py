import openpyxl
from docx import Document
import sys
import os

def analyze_excel(file_path):
    print(f"--- Analyzing Excel: {os.path.basename(file_path)} ---")
    try:
        wb = openpyxl.load_workbook(file_path)
        ws = wb.active
        # Read first row
        headers = [cell.value for cell in ws[1]]
        print(f"Total Columns: {len(headers)}")
        print("Columns:", headers)
    except Exception as e:
        print(f"Error reading Excel: {e}")

def analyze_docx(file_path):
    print(f"\n--- Analyzing Word: {os.path.basename(file_path)} ---")
    try:
        doc = Document(file_path)
        print(f"Total Paragraphs: {len(doc.paragraphs)}")
        print("First 10 Paragraphs:")
        for p in doc.paragraphs[:10]:
            if p.text.strip():
                print(f"- {p.text.strip()}")
        
        print("\nTables found:", len(doc.tables))
        if doc.tables:
            print("First table structure (first row):")
            for cell in doc.tables[0].rows[0].cells:
                print(f"[{cell.text.strip()}]", end=" ")
            print()
    except Exception as e:
        print(f"Error reading Word: {e}")

if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python analyze_templates.py <excel_path> <docx_path>")
    else:
        analyze_excel(sys.argv[1])
        analyze_docx(sys.argv[2])
